#!/usr/bin/env python3
"""Генерация документации Merchant API в формате Word (.docx).

Зависимость: python-docx (НЕ пакет «docx» — он другой и ломает import).

  pip uninstall docx -y
  pip install -r docs/requirements.txt
"""

from pathlib import Path
import sys

try:
    from docx import Document
except ModuleNotFoundError as exc:
    if exc.name == "exceptions":
        sys.exit(
            "Ошибка: установлен неверный пакет «docx».\n"
            "  pip uninstall docx -y\n"
            "  pip install python-docx\n"
            "или: pip install -r docs/requirements.txt"
        )
    raise
except ImportError:
    sys.exit(
        "Не установлен python-docx.\n"
        "  pip install -r docs/requirements.txt"
    )
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn


OUTPUT = Path(__file__).resolve().parent / "AvaPay_Merchant_API.docx"


def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def add_title(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_number(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    pf = p.paragraph_format
    pf.left_indent = Cm(0.5)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def build() -> None:
    doc = Document()
    set_default_font(doc)

    # Титул
    add_title(doc, "AvaPay — документация Merchant API")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Версия API: v1  |  Формат: JSON  |  Кодировка: UTF-8").italic = True
    doc.add_paragraph()

    add_para(
        doc,
        "Руководство для мерчантов AvaPay: приём платежей (pay-in), выплаты (pay-out), "
        "статусы операций, callback-уведомления и апелляции.",
    )
    doc.add_page_break()

    # 1. Обзор
    add_heading(doc, "1. Обзор")
    add_para(doc, "Merchant API позволяет:")
    for item in [
        "Создавать заявки на входящие платежи и получать реквизиты для оплаты",
        "Создавать заявки на исходящие выплаты",
        "Получать callback при смене статуса операции",
        "Управлять API-ключами и IP whitelist",
        "Подавать апелляции по спорным pay-in транзакциям",
    ]:
        add_bullet(doc, item)
    add_para(doc, "Два режима pay-in:", bold=True)
    add_bullet(doc, "H2H — реквизиты сразу в ответе API (поле payment_details)")
    add_bullet(doc, "Invoice — ссылка на платёжную страницу (поле redirect_url)")

    # 2. Базовый URL
    add_heading(doc, "2. Базовый URL")
    add_code(doc, "https://api.avapay.net/api/v1/payments/")
    add_para(doc, "Ключи, валюты и платёжные системы доступны в личном кабинете мерчанта (раздел API Keys).")

    # 3. Аутентификация
    add_heading(doc, "3. Аутентификация")
    add_para(doc, "Каждый запрос к API должен содержать заголовок:")
    add_code(doc, "Authorization: Token <ваш_api_token>")
    add_table(
        doc,
        ["Параметр", "Описание"],
        [
            ["Token", "API token, выдаётся при создании пары ключей"],
            ["private_key", "UUID для подписи запросов и проверки callback. Храните только на сервере."],
        ],
    )
    add_para(
        doc,
        "ВАЖНО: не передавайте private_key во frontend, мобильные приложения и публичную документацию.",
        bold=True,
    )

    # 4. Подпись
    add_heading(doc, "4. Подпись запросов")
    add_para(
        doc,
        "Для всех POST-запросов с телом (pay-in, pay-out, whitelist) обязателен заголовок Signature.",
    )
    add_heading(doc, "4.1. Алгоритм", level=2)
    add_number(doc, "Сериализуйте JSON-тело запроса с сортировкой ключей, без пробелов")
    add_number(doc, "Добавьте к строке ваш private_key (как текст)")
    add_number(doc, "Вычислите SHA-256 hex-дайджест")
    add_code(doc, "Signature = SHA256( JSON_sorted(body) + private_key )")
    add_heading(doc, "4.2. Пример (Python)", level=2)
    add_code(
        doc,
        """import hashlib
import json

PRIVATE_KEY = "00000000-0000-4000-8000-000000000001"

request_data = {
    "amount": "2000.00",
    "currency": "KZT",
    "payment_system": "C2CKZT",
    "merchant_order_id": "order-demo-10001",
    "callback_url": "https://merchant.example.com/callbacks/avapay",
    "ftd": False,
    "client": {
        "client_id": "user-demo-42",
        "email": "client@example.com",
        "phone": "+77000000000",
        "name": "Demo Client"
    }
}

sorted_json = json.dumps(request_data, sort_keys=True, separators=(",", ":"))
signature = hashlib.sha256((sorted_json + PRIVATE_KEY).encode()).hexdigest()

headers = {
    "Authorization": "Token <token>",
    "Content-Type": "application/json",
    "Signature": signature,
}""",
    )
    add_para(doc, "Подписывайте точно то тело, которое отправляете в запросе.")

    add_heading(doc, "4.3. Подпись ответов API", level=2)
    add_para(
        doc,
        "При успешном создании заявки (HTTP 201) и при GET-запросе статуса AvaPay возвращает "
        "заголовок Signature — подпись JSON-тела ответа тем же алгоритмом. Рекомендуется проверять её.",
    )

    # 5. Callback
    add_heading(doc, "5. Callback-уведомления")
    add_para(
        doc,
        "При смене статуса pay-in / pay-out AvaPay отправляет POST на callback_url, "
        "указанный при создании заявки.",
    )
    add_heading(doc, "5.1. Тело callback", level=2)
    add_code(
        doc,
        """{
  "id": "a1b2c3d4-e5f6-7890-abcd-ef1234567890",
  "order_id": "order-demo-10001",
  "amount": 2000.0,
  "currency": "KZT",
  "payment_system": "C2CKZT",
  "status": "Success",
  "recalculated": false,
  "timestamp": 1718123456
}""",
    )
    add_para(doc, "Поле order_id — это ваш merchant_order_id.")
    add_heading(doc, "5.2. Заголовок", level=2)
    add_code(doc, "Signature: <sha256 подпись тела callback + private_key>")
    add_para(doc, "Ответьте HTTP 2xx. При ошибке доставки возможны повторные попытки.")

    # 6. API Keys
    add_heading(doc, "6. API Keys")
    add_heading(doc, "6.1. Создание ключа", level=2)
    add_para(doc, "POST /api/v1/payments/keys/")
    add_para(doc, "Тело запроса: пустой объект {}")
    add_para(doc, "Создаёт новую активную пару ключей. Предыдущие активные ключи деактивируются.")
    add_heading(doc, "6.2. Ответ 201", level=2)
    add_code(
        doc,
        """{
  "id": "9e3e624c-175c-496a-a7e7-e408ed671163",
  "token": "...",
  "private_key": "...",
  "created_at": "2026-06-01T12:00:00Z",
  "whitelist_on": false,
  "whitelist_ips": []
}""",
    )
    add_para(doc, "private_key отображается только при создании. Сохраните его сразу.")
    add_heading(doc, "6.3. Список ключей", level=2)
    add_para(doc, "GET /api/v1/payments/keys/")
    add_heading(doc, "6.4. IP whitelist", level=2)
    add_para(doc, "POST /api/v1/payments/keys/{id}/whitelist/")
    add_code(
        doc,
        """{
  "whitelist_on": true,
  "whitelist": ["203.0.113.10", "203.0.113.11"]
}""",
    )

    doc.add_page_break()

    # 7. Pay-in H2H
    add_heading(doc, "7. Pay-in — H2H")
    add_para(doc, "POST /api/v1/payments/in/h2h/")
    add_para(doc, "Создание заявки с немедленной выдачей реквизитов в поле payment_details.")
    add_heading(doc, "7.1. Тело запроса", level=2)
    add_code(
        doc,
        """{
  "amount": "3000.00",
  "currency": "KZT",
  "payment_system": "C2CKZT",
  "merchant_order_id": "order-demo-10002",
  "callback_url": "https://merchant.example.com/callbacks/avapay",
  "success_url": "https://merchant.example.com/payment/success",
  "failed_url": "https://merchant.example.com/payment/failed",
  "ftd": false,
  "client": {
    "client_id": "user-demo-42",
    "email": "client@example.com",
    "phone": "+77000000000",
    "name": "Demo Client"
  }
}""",
    )
    add_table(
        doc,
        ["Поле", "Обяз.", "Описание"],
        [
            ["amount", "да", "Сумма в валюте операции"],
            ["currency", "да", "Код валюты: KZT, RUB и др."],
            ["payment_system", "да", "Sber, SBP, C2C, C2CKZT и др."],
            ["merchant_order_id", "да", "Уникальный ID заказа в вашей системе"],
            ["callback_url", "да", "URL для webhook статусов"],
            ["success_url", "нет", "Редирект при успехе (invoice-режим)"],
            ["failed_url", "нет", "Редирект при ошибке"],
            ["ftd", "да", "true — первый депозит, false — повторный"],
            ["client", "да", "Объект клиента"],
        ],
    )
    add_table(
        doc,
        ["client.*", "Обяз.", "Описание"],
        [
            ["client_id", "да", "Стабильный ID пользователя у мерчанта"],
            ["email", "нет", "Email клиента"],
            ["phone", "нет", "Телефон в международном формате"],
            ["name", "нет", "Имя плательщика"],
        ],
    )
    add_heading(doc, "7.2. Ответ 201 (успех)", level=2)
    add_code(
        doc,
        """{
  "id": "a1b2c3d4-e5f6-7890-abcd-ef1234567890",
  "currency": "KZT",
  "amount": "3000.00",
  "payment_system": "C2CKZT",
  "status": "New",
  "merchant_order_id": "order-demo-10002",
  "payment_details": {
    "card_number": "4000000000000001",
    "owner": "Иванов И. И.",
    "bank": "Example Bank"
  },
  "expires_at": 1718127056,
  "recalculated": false,
  "usd_amount": 5.77
}""",
    )
    add_para(doc, "Формат payment_details по payment_system:")
    add_bullet(doc, "Sber, C2C, C2CKZT — card_number, owner, bank")
    add_bullet(doc, "SBP, SberPay — phone, owner, bank")
    add_bullet(doc, "SberDep — deposit_number, bic, owner")

    add_heading(doc, "7.3. Отказ при создании (HTTP 400)", level=2)
    add_para(
        doc,
        "Если реквизиты не удалось выдать, API возвращает HTTP 400 (без callback Declined):",
    )
    add_code(
        doc,
        """{
  "error": "Не удалось выдать платёжные реквизиты для указанной суммы и метода оплаты.",
  "error_code": "routing_unavailable",
  "pay_in_id": "...",
  "in_order_id": "..."
}""",
    )
    add_table(
        doc,
        ["error_code", "Описание"],
        [
            ["routing_unavailable", "Роутинг не нашёл реквизиты / все PSP отказали"],
            ["requisites_unavailable", "Временно нет реквизитов — повторите позже"],
            ["requisites_empty_response", "PSP вернул пустой ответ без реквизитов"],
        ],
    )

    # 8. Pay-in Invoice
    add_heading(doc, "8. Pay-in — Invoice (редирект)")
    add_para(doc, "POST /api/v1/payments/in/invoice/")
    add_para(doc, "Те же поля, что у H2H. В ответе дополнительно redirect_url — ссылка на hosted-страницу оплаты.")
    add_para(doc, "Начальный статус: In Progress. Пример redirect_url: https://pay.avapay.net/{pay_in_id}")
    add_para(doc, "Схема: POST invoice → redirect_url → клиент оплачивает → callback_url → success_url / failed_url")
    add_heading(doc, "8.1. Публичное API страницы", level=2)
    add_table(
        doc,
        ["Метод", "URL", "Описание"],
        [
            ["GET", "/api/v1/payments/in/invoice/{id}/obtain/", "Реквизиты, таймер, locale (kk/ru)"],
            ["POST", ".../sent/", "Клиент нажал «Я оплатил»"],
            ["POST", ".../cancel/", "Отмена"],
            ["GET", ".../complete/", "Терминальный статус + URL редиректа"],
        ],
    )
    add_heading(doc, "8.2. Deep links (Kaspi / Homebank)", level=2)
    add_para(
        doc,
        "Кнопки открывают приложение банка и копируют реквизиты («Скопировать всё»). "
        "Автоподстановка карты и суммы внутри Kaspi без merchant API Kaspi Pay невозможна — "
        "клиент вставляет данные вручную (Перевод → на карту другого банка).",
    )
    add_heading(doc, "8.3. Развёртывание (администратор)", level=2)
    add_para(doc, "Файлы: templates/payment_page/pay.html, payments/payment_page.py, payment_page_enrich.py, bank_deeplinks.py, invoice_obtain.py, viewsets.py, urls.py, trade/utils2.py")
    add_para(doc, "Traefik: pay.avapay.net → app:8080, PAYMENT_PAGE_URL=avapay.net, DNS A для pay/payment/payments")
    add_para(doc, "HTTPS: Let's Encrypt через Traefik certresolver. Favicon — inline SVG в pay.html")

    # 9. Статусы
    add_heading(doc, "9. Статусы pay-in")
    add_table(
        doc,
        ["Статус", "Описание"],
        [
            ["New", "Заявка создана, ожидается оплата (H2H)"],
            ["In Progress", "В обработке (invoice-режим)"],
            ["Success", "Платёж подтверждён"],
            ["Failed", "Ошибка / истечение без оплаты"],
            ["Declined", "Отклонена при создании (нет реквизитов, лимиты)"],
        ],
    )
    add_para(doc, "expires_at — Unix timestamp, до которого действительны реквизиты.")

    # 10. Доп. действия
    add_heading(doc, "10. Дополнительные действия pay-in")
    add_table(
        doc,
        ["Метод", "URL", "Описание"],
        [
            ["GET", "/api/v1/payments/in/h2h/{id}/", "Получить заявку"],
            ["POST", "/api/v1/payments/in/h2h/{id}/sent/", "Отметить «деньги отправлены»"],
            ["POST", "/api/v1/payments/in/h2h/{id}/cancel/", "Отмена заявки"],
            ["POST", "/api/v1/payments/in/h2h/{id}/arbitrage/", "Апелляция (multipart, file: PNG/JPEG/PDF до 5 MB)"],
        ],
    )

    doc.add_page_break()

    # 11. Pay-out
    add_heading(doc, "11. Pay-out — H2H")
    add_para(doc, "POST /api/v1/payments/out/h2h/")
    add_code(
        doc,
        """{
  "amount": "5000.00",
  "currency": "RUB",
  "payment_system": "Sber",
  "merchant_order_id": "payout-demo-20001",
  "callback_url": "https://merchant.example.com/callbacks/avapay-payout",
  "ftd": false,
  "client": { "client_id": "user-demo-42", ... },
  "details": { "card_number": "4000000000000002" }
}""",
    )
    add_para(doc, "Поле details обязательно. Набор полей — см. required_fields в справочнике payment-systems.")

    # 12. Справочники
    add_heading(doc, "12. Справочники")
    add_heading(doc, "12.1. Валюты", level=2)
    add_para(doc, "GET /api/v1/payments/currencies/")
    add_heading(doc, "12.2. Платёжные системы", level=2)
    add_para(doc, "GET /api/v1/payments/payment-systems/")

    # 13. Ошибки
    add_heading(doc, "13. Типичные ошибки")
    add_table(
        doc,
        ["HTTP", "Сообщение", "Причина"],
        [
            ["403", "Invalid signature", "Неверная подпись тела запроса"],
            ["403", "Signature required", "Нет заголовка Signature на POST"],
            ["403", "This IP is not in whitelist", "IP не в whitelist API-ключа"],
            ["400", "This method is not active", "Метод не подключён или неверный ftd"],
            ["400", "Amount out of limits!", "Сумма вне лимитов для метода"],
            ["400", "Order with such merchant_order_id already exists", "Дубликат ID заказа"],
            ["400", "error + error_code", "Не удалось выдать реквизиты (см. раздел 7.3)"],
        ],
    )

    # 14. Безопасность
    add_heading(doc, "14. Безопасность")
    for item in [
        "Храните private_key только на backend",
        "Используйте HTTPS для callback URL",
        "Включите IP whitelist в production",
        "Проверяйте подпись каждого callback и ответа API",
        "Не логируйте полные реквизиты карт в открытом виде",
        "Ротируйте ключи при компрометации",
    ]:
        add_bullet(doc, item)

    # 15. Чеклист
    add_heading(doc, "15. Чеклист запуска")
    for i, item in enumerate(
        [
            "Получите доступ в личный кабинет AvaPay",
            "Создайте API Keys, сохраните token и private_key",
            "Реализуйте подпись исходящих POST-запросов",
            "Реализуйте приём и проверку callback",
            "Проверьте доступные payment_system и лимиты для вашего аккаунта",
            "Проведите тестовые заявки на минимальных суммах",
            "Включите IP whitelist перед production",
        ],
        start=1,
    ):
        add_number(doc, item)

    add_para(doc, "")
    add_para(doc, "Техническая поддержка: обращайтесь к вашему менеджеру AvaPay.")

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build()
